from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, replace
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

TZ = ZoneInfo("Asia/Tokyo")

LABELS = {
    "ja": {
        "report": "専門調査レポート", "status": "検証状態", "validated": "検証済み", "needs_review": "要確認",
        "summary": "エグゼクティブサマリー", "scope": "調査範囲・方法", "findings": "主要な調査結果",
        "analysis": "専門的分析・示唆", "limitations": "リスク・制約・未確認事項", "recommendations": "推奨アクション",
        "sources": "出典", "no_items": "該当なし",
        "method": "本レポートは、Agent 1 が検証した根拠のみを使用し、確認済み事実・推論・未確認事項を区別して作成しています。",
        "audit_note": "技術ログ、内部Evidence ID、ワークフロー履歴は本文から分離し、監査用成果物として保存しています。",
    },
    "vi": {
        "report": "Báo cáo phân tích chuyên sâu", "status": "Trạng thái xác minh", "validated": "Đã xác minh", "needs_review": "Cần xác nhận",
        "summary": "Tóm tắt điều hành", "scope": "Phạm vi và phương pháp", "findings": "Các phát hiện chính",
        "analysis": "Phân tích chuyên gia và hàm ý", "limitations": "Rủi ro, giới hạn và nội dung chưa xác minh", "recommendations": "Khuyến nghị hành động",
        "sources": "Nguồn tham khảo", "no_items": "Không có",
        "method": "Báo cáo chỉ sử dụng bằng chứng đã được Agent 1 xác minh, đồng thời phân biệt rõ dữ kiện, suy luận và nội dung chưa được xác nhận.",
        "audit_note": "Log kỹ thuật, Evidence ID nội bộ và lịch sử workflow được tách khỏi báo cáo chính và lưu riêng để phục vụ kiểm toán.",
    },
    "en": {
        "report": "Expert Research Report", "status": "Validation status", "validated": "Validated", "needs_review": "Needs review",
        "summary": "Executive Summary", "scope": "Scope and Method", "findings": "Key Findings",
        "analysis": "Expert Analysis and Implications", "limitations": "Risks, Limitations and Unresolved Items", "recommendations": "Recommended Actions",
        "sources": "Sources", "no_items": "None",
        "method": "This report uses only evidence validated by Agent 1 and separates verified facts, inferences and unresolved items.",
        "audit_note": "Technical logs, internal evidence IDs and workflow history are separated from this reader-facing report and preserved as audit artifacts.",
    },
}
LANGUAGE_NAMES = {"ja": "Japanese", "vi": "Vietnamese", "en": "English"}
ENGLISH_STOPWORDS = {"the", "and", "that", "this", "with", "from", "for", "are", "is", "was", "were", "have", "has", "will", "should", "based", "source", "sources", "analysis", "report", "research", "specific", "information", "technical", "recommended", "use", "using"}
VI_MARKERS = {"và", "của", "các", "được", "trong", "với", "cho", "là", "những", "này", "cần", "không", "nghiên", "cứu", "phân", "tích", "khuyến", "nghị", "nguồn", "xác", "minh"}


@dataclass(frozen=True)
class HumanReportBundle:
    markdown: str
    paths: tuple[str, ...]
    warnings: tuple[str, ...] = ()


def _lang(language: str) -> str:
    value = (language or "ja").strip().lower()
    return value if value in LABELS else "ja"


def _clean(value: Any) -> str:
    return " ".join(str(value or "").split())


def _refs(item: dict[str, Any]) -> str:
    ids = item.get("source_ids") if isinstance(item, dict) else []
    refs = [str(x).strip() for x in ids] if isinstance(ids, list) else []
    refs = [x for x in refs if x]
    return " " + " ".join(f"[{x}]" for x in refs) if refs else ""


def _valid_source_ids(handoff: dict[str, Any]) -> set[str]:
    return {_clean(x.get("source_id")) for x in handoff.get("sources", []) if isinstance(x, dict) and _clean(x.get("source_id"))}


def _valid_refs(value: Any, valid: set[str]) -> list[str]:
    if not isinstance(value, list):
        return []
    return [sid for sid in (_clean(x) for x in value) if sid in valid]


def build_report_data(task_id: str, title: str, request: str, handoff: dict[str, Any], language: str) -> dict[str, Any]:
    """Build an evidence-preserving deterministic fallback for offline callers."""
    language = _lang(language)
    labels = LABELS[language]
    facts = [x for x in handoff.get("key_facts", []) if isinstance(x, dict) and _clean(x.get("claim"))]
    inferences = [x for x in handoff.get("inferences", []) if isinstance(x, dict) and _clean(x.get("claim"))]
    limitations: list[str] = []
    for conflict in handoff.get("conflicts", []):
        if not isinstance(conflict, dict):
            continue
        topic, description = _clean(conflict.get("topic")), _clean(conflict.get("description"))
        if topic or description:
            limitations.append(f"{topic}: {description}" if topic and description else topic or description)
    limitations.extend(_clean(x) for x in handoff.get("unresolved_items", []) if _clean(x))
    if handoff.get("presentation_ready") is not True:
        limitations = [_clean(x) for x in handoff.get("blockers", []) if _clean(x)] + limitations
    actions = [_clean(x) for x in handoff.get("recommended_next_actions", []) if _clean(x)]
    return {
        "task_id": task_id,
        "title": _clean(title) or task_id,
        "request": _clean(request),
        "language": language,
        "report_label": labels["report"],
        "status_label": labels["status"],
        "status": labels["validated"] if handoff.get("presentation_ready") is True else labels["needs_review"],
        "summary": _clean(handoff.get("conclusion")) or labels["method"],
        "scope": labels["method"],
        "facts": facts,
        "analysis_items": [{"text": _clean(x.get("claim")), "source_ids": list(x.get("source_ids") or [])} for x in inferences],
        "limitations": list(dict.fromkeys(limitations)),
        "actions": [{"priority": "P1", "text": x, "source_ids": []} for x in list(dict.fromkeys(actions))],
        "sources": [x for x in handoff.get("sources", []) if isinstance(x, dict) and _clean(x.get("url"))],
        "labels": labels,
        "generated_at": datetime.now(TZ).isoformat(),
        "expert_composed": False,
    }


def _model_payload(data: dict[str, Any], handoff: dict[str, Any]) -> dict[str, Any]:
    return {
        "task_id": data["task_id"],
        "original_title": data["title"],
        "user_request": data["request"],
        "target_language": LANGUAGE_NAMES[data["language"]],
        "presentation_ready": handoff.get("presentation_ready") is True,
        "blockers": handoff.get("blockers", []),
        "verified_facts": handoff.get("key_facts", []),
        "inferences": handoff.get("inferences", []),
        "conflicts": handoff.get("conflicts", []),
        "unresolved_items": handoff.get("unresolved_items", []),
        "conclusion": handoff.get("conclusion", ""),
        "recommended_next_actions": handoff.get("recommended_next_actions", []),
        "sources": [{"source_id": x.get("source_id"), "title": x.get("title"), "url": x.get("url")} for x in handoff.get("sources", []) if isinstance(x, dict)],
    }


def _language_ok(text: str, language: str) -> bool:
    language = _lang(language)
    body = re.sub(r"https?://\S+|\[[A-Z]\d+\]|`[^`]*`", " ", text)
    if language == "en":
        return True
    if language == "ja":
        meaningful = re.sub(r"[\s\d\W_]+", "", body)
        japanese = len(re.findall(r"[\u3040-\u30ff\u3400-\u9fff]", meaningful))
        latin_words = re.findall(r"[A-Za-z]{3,}", body)
        return japanese >= 40 and japanese >= len(latin_words) * 2
    words = re.findall(r"[A-Za-zÀ-ỹĐđ]+", body.lower())
    vi_hits = sum(word in VI_MARKERS for word in words)
    en_hits = sum(word in ENGLISH_STOPWORDS for word in words)
    diacritics = len(re.findall(r"[À-ỹĐđ]", body))
    return bool(words) and (vi_hits >= 5 or diacritics >= 20) and en_hits <= max(8, vi_hits)


def _language_body(data: dict[str, Any]) -> str:
    values = [data.get("title", ""), data.get("summary", ""), data.get("scope", "")]
    for item in data.get("facts", []):
        values.extend((_clean(item.get("heading")), _clean(item.get("claim"))))
    for item in data.get("analysis_items", []):
        values.extend((_clean(item.get("heading")), _clean(item.get("text"))))
    values.extend(_clean(x) for x in data.get("limitations", []))
    for item in data.get("actions", []):
        values.append(_clean(item.get("text") if isinstance(item, dict) else item))
    return "\n".join(str(x) for x in values if x)


def _report_items(value: Any, valid: set[str], *, refs_required: bool, max_items: int) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    result: list[dict[str, Any]] = []
    for raw in value:
        if not isinstance(raw, dict):
            continue
        text = _clean(raw.get("text"))
        refs = _valid_refs(raw.get("source_ids"), valid)
        if not text or (refs_required and not refs):
            continue
        item: dict[str, Any] = {"text": text, "source_ids": refs}
        heading = _clean(raw.get("heading"))
        priority = _clean(raw.get("priority")).upper()
        if heading:
            item["heading"] = heading
        if priority in {"P0", "P1", "P2"}:
            item["priority"] = priority
        result.append(item)
        if len(result) >= max_items:
            break
    return result


def _normalize_expert(raw: dict[str, Any], base: dict[str, Any], handoff: dict[str, Any]) -> dict[str, Any]:
    valid = _valid_source_ids(handoff)
    findings = _report_items(raw.get("findings"), valid, refs_required=True, max_items=12)
    analysis = _report_items(raw.get("analysis"), valid, refs_required=True, max_items=10)
    actions = _report_items(raw.get("recommendations"), valid, refs_required=False, max_items=8)
    limitations = [_clean(x) for x in raw.get("limitations", []) if _clean(x)] if isinstance(raw.get("limitations"), list) else []
    if not findings:
        findings = [{"text": _clean(x.get("claim")), "source_ids": _valid_refs(x.get("source_ids"), valid)} for x in handoff.get("key_facts", []) if isinstance(x, dict) and _clean(x.get("claim")) and _valid_refs(x.get("source_ids"), valid)][:12]
    data = dict(base)
    data["title"] = _clean(raw.get("report_title")) or base["title"]
    data["summary"] = _clean(raw.get("executive_summary")) or base["summary"]
    data["scope"] = _clean(raw.get("scope_and_method")) or base["scope"]
    data["facts"] = [{"claim": x["text"], "source_ids": x["source_ids"], **({"heading": x["heading"]} if x.get("heading") else {})} for x in findings]
    data["analysis_items"] = analysis
    data["limitations"] = limitations or base["limitations"]
    data["actions"] = actions or base["actions"]
    data["expert_composed"] = True
    return data


def _expert_prompt(payload: dict[str, Any], repair: bool = False) -> str:
    target = payload["target_language"]
    retry = "The previous draft failed the target-language gate. Rewrite it completely.\n" if repair else ""
    return f"""{retry}Create a reader-facing expert research report.
TARGET LANGUAGE: {target}

NON-NEGOTIABLE LANGUAGE RULE
- Write ALL reader-facing narrative in {target}: title, summary, method, findings, analysis, limitations and recommendations.
- For Japanese/Vietnamese output, do not leave English prose in the report.
- Proper nouns, product names, exact technical identifiers, source titles, URLs, code, standards and S1/S2 source IDs may remain unchanged.

PROFESSIONAL STANDARD
- Write as a senior R&D/technology analyst for managers and engineers.
- Use formal, precise, domain-appropriate language. No casual chat, slogans, filler or shallow one-line commentary.
- Executive summary must synthesize the decision-relevant conclusion in 2-4 substantive paragraphs when evidence permits.
- Findings must explain what is established and why it matters.
- Analysis must explain implications, dependencies, trade-offs and operational relevance without exceeding the evidence.
- Recommendations must be concrete and prioritized P0/P1/P2 where justified.
- State weak evidence and uncertainty explicitly.

EVIDENCE BOUNDARY
- Use ONLY the JSON evidence supplied below. Never invent facts, dates, numbers, causes, products or market claims.
- Every finding and analytical point must cite valid source_ids.
- Recommendations may have empty source_ids only when clearly presented as proposed next actions.
- Preserve conflicts and unresolved items. Do not rewrite source titles or URLs.

Return JSON only:
{{
  "report_title":"...",
  "executive_summary":"...",
  "scope_and_method":"...",
  "findings":[{{"heading":"...","text":"...","source_ids":["S1"]}}],
  "analysis":[{{"heading":"...","text":"...","source_ids":["S1"]}}],
  "limitations":["..."],
  "recommendations":[{{"priority":"P0|P1|P2","text":"...","source_ids":["S1"]}}]
}}

EVIDENCE:\n{json.dumps(payload, ensure_ascii=False, indent=2)}"""


def compose_expert_report(data: dict[str, Any], handoff: dict[str, Any], llm: Any) -> dict[str, Any]:
    payload = _model_payload(data, handoff)
    reason = ""
    for attempt in range(2):
        raw = llm.generate_json(
            "You are the 3Agent Expert Report Composer. Follow language and evidence constraints exactly.",
            _expert_prompt(payload, repair=attempt > 0),
            think=False,
            num_predict=6144,
        )
        candidate = _normalize_expert(raw, data, handoff)
        if _language_ok(_language_body(candidate), data["language"]):
            return candidate
        reason = f"target-language validation failed for {LANGUAGE_NAMES[data['language']]} on attempt {attempt + 1}"
        payload["previous_invalid_draft"] = raw
    raise ValueError(f"Human report rejected: {reason}. Refusing to publish a mixed-language report.")


def _runtime_report_llm() -> Any | None:
    """Resolve the configured local report model while preserving resource admission."""
    config_path = os.getenv("THREE_AGENT_CONFIG", "").strip()
    if not config_path or not Path(config_path).is_file():
        return None
    try:
        from .config import load_config, legacy_model_policy
        from .llm import OllamaClient
        from .resource_budget import ResourceBudgetConfig, ResourceBudgetManager

        config = load_config()
        policy = config.model_policy or legacy_model_policy(config.llm)
        manager = None
        if policy.enabled and policy.resource_control_enabled:
            manager = ResourceBudgetManager(
                config.llm.base_url,
                ResourceBudgetConfig(
                    enabled=True,
                    max_vram_percent=policy.max_vram_percent,
                    max_ram_percent=policy.max_ram_percent,
                    max_gpu_util_percent=policy.max_gpu_util_percent,
                    max_gpu_power_percent=policy.max_gpu_power_percent,
                    max_gpu_temp_c=policy.max_gpu_temp_c,
                    model_size_safety_factor=policy.model_size_safety_factor,
                    model_ram_overhead_factor=policy.model_ram_overhead_factor,
                    serialize_generation=policy.serialize_generation,
                    reservation_ttl_seconds=policy.reservation_ttl_seconds,
                ),
            )
        model = policy.report_model if policy.enabled else config.llm.model
        return OllamaClient(replace(config.llm, model=model), manager)
    except Exception:
        return None


def render_markdown(data: dict[str, Any]) -> str:
    labels = data["labels"]
    lines = [f"# {data['title']}", "", f"> {data['report_label']} · {data['status_label']}: **{data['status']}**", "", f"## {labels['summary']}", data["summary"], "", f"## {labels['scope']}", data.get("scope") or labels["method"], "", f"## {labels['findings']}"]
    if data["facts"]:
        for index, item in enumerate(data["facts"], 1):
            heading, claim = _clean(item.get("heading")), _clean(item.get("claim") or item.get("text"))
            if heading:
                lines.extend([f"### {index}. {heading}", f"{claim}{_refs(item)}"])
            else:
                lines.append(f"- {claim}{_refs(item)}")
    else:
        lines.append(f"- {labels['no_items']}")
    lines.extend(["", f"## {labels['analysis']}"])
    analysis = data.get("analysis_items") or []
    if analysis:
        for index, item in enumerate(analysis, 1):
            heading, text = _clean(item.get("heading")), _clean(item.get("text"))
            if heading:
                lines.extend([f"### {index}. {heading}", f"{text}{_refs(item)}"])
            else:
                lines.append(f"- {text}{_refs(item)}")
    else:
        lines.append(f"- {labels['no_items']}")
    lines.extend(["", f"## {labels['limitations']}"])
    lines.extend(f"- {x}" for x in data["limitations"] or [labels["no_items"]])
    lines.extend(["", f"## {labels['recommendations']}"])
    actions = data.get("actions") or []
    if actions:
        for item in actions:
            if isinstance(item, dict):
                priority = _clean(item.get("priority"))
                lines.append(f"- {f'**{priority}** — ' if priority else ''}{_clean(item.get('text'))}{_refs(item)}")
            else:
                lines.append(f"- {_clean(item)}")
    else:
        lines.append(f"- {labels['no_items']}")
    lines.extend(["", f"## {labels['sources']}"])
    if data["sources"]:
        for source in data["sources"]:
            sid, name = _clean(source.get("source_id")) or "S?", _clean(source.get("title"))
            lines.append(f"- [{sid}] {name or sid} — {_clean(source.get('url'))}")
    else:
        lines.append(f"- {labels['no_items']}")
    lines.extend(["", "---", "", labels["audit_note"], "", f"Task: `{data['task_id']}`"])
    return "\n".join(lines).rstrip() + "\n"


def _set_docx_font(run: Any, language: str, size: int | None = None, bold: bool | None = None) -> None:
    font = "Noto Sans CJK JP" if language == "ja" else "Aptos"
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _write_docx(data: dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Cm(1.8), Cm(1.8), Cm(2.0), Cm(2.0)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _set_docx_font(p.add_run(data["title"]), data["language"], 20, True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _set_docx_font(p.add_run(f"{data['report_label']} · {data['status_label']}: {data['status']}"), data["language"], 10)
    def heading(text: str, size: int = 14) -> None:
        p = doc.add_paragraph(); _set_docx_font(p.add_run(text), data["language"], size, True)
    def paragraph(text: str) -> None:
        p = doc.add_paragraph(); _set_docx_font(p.add_run(text), data["language"], 10)
    def section_items(label: str, items: list[dict[str, Any]], text_key: str) -> None:
        heading(label)
        if not items:
            paragraph(data["labels"]["no_items"]); return
        for index, item in enumerate(items, 1):
            sub = _clean(item.get("heading"))
            if sub:
                heading(f"{index}. {sub}", 11)
            paragraph(_clean(item.get(text_key)) + _refs(item))
    heading(data["labels"]["summary"]); paragraph(data["summary"])
    heading(data["labels"]["scope"]); paragraph(data.get("scope") or data["labels"]["method"])
    section_items(data["labels"]["findings"], data["facts"], "claim")
    section_items(data["labels"]["analysis"], data.get("analysis_items") or [], "text")
    heading(data["labels"]["limitations"])
    for text in data["limitations"] or [data["labels"]["no_items"]]: paragraph("• " + text)
    heading(data["labels"]["recommendations"])
    for item in data.get("actions") or []:
        if isinstance(item, dict): paragraph(f"• {_clean(item.get('priority'))} — {_clean(item.get('text'))}{_refs(item)}")
        else: paragraph("• " + _clean(item))
    heading(data["labels"]["sources"])
    for source in data["sources"]:
        paragraph(f"[{_clean(source.get('source_id')) or 'S?'}] {_clean(source.get('title'))} — {_clean(source.get('url'))}")
    paragraph(data["labels"]["audit_note"]); paragraph(f"Task: {data['task_id']}")
    doc.save(path)


def _pdf_font(language: str) -> str:
    if language == "ja":
        name = "HeiseiKakuGo-W5"
        try: pdfmetrics.getFont(name)
        except KeyError: pdfmetrics.registerFont(UnicodeCIDFont(name))
        return name
    regular = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
    if regular.is_file():
        name = "ThreeAgentDejaVu"
        try: pdfmetrics.getFont(name)
        except KeyError: pdfmetrics.registerFont(TTFont(name, str(regular)))
        return name
    return "Helvetica"


def _write_pdf(data: dict[str, Any], path: Path) -> None:
    font = _pdf_font(data["language"]); styles = getSampleStyleSheet()
    title = ParagraphStyle("TA_Title", parent=styles["Title"], fontName=font, fontSize=18, leading=23, spaceAfter=8)
    meta = ParagraphStyle("TA_Meta", parent=styles["Normal"], fontName=font, fontSize=9, leading=12, textColor="#555555", spaceAfter=10)
    head = ParagraphStyle("TA_Head", parent=styles["Heading2"], fontName=font, fontSize=13, leading=17, spaceBefore=10, spaceAfter=5)
    body = ParagraphStyle("TA_Body", parent=styles["BodyText"], fontName=font, fontSize=9.5, leading=14, spaceAfter=4)
    story = [Paragraph(escape(data["title"]), title), Paragraph(escape(f"{data['report_label']} · {data['status_label']}: {data['status']}"), meta)]
    def add_heading(label: str) -> None: story.append(Paragraph(escape(label), head))
    def add_text(text: str) -> None: story.append(Paragraph(escape(text), body))
    add_heading(data["labels"]["summary"]); add_text(data["summary"])
    add_heading(data["labels"]["scope"]); add_text(data.get("scope") or data["labels"]["method"])
    add_heading(data["labels"]["findings"])
    for index, item in enumerate(data["facts"], 1):
        if _clean(item.get("heading")): add_text(f"{index}. {_clean(item.get('heading'))}")
        add_text(_clean(item.get("claim")) + _refs(item))
    add_heading(data["labels"]["analysis"])
    for index, item in enumerate(data.get("analysis_items") or [], 1):
        if _clean(item.get("heading")): add_text(f"{index}. {_clean(item.get('heading'))}")
        add_text(_clean(item.get("text")) + _refs(item))
    add_heading(data["labels"]["limitations"])
    for text in data["limitations"] or [data["labels"]["no_items"]]: add_text("• " + text)
    add_heading(data["labels"]["recommendations"])
    for item in data.get("actions") or []:
        if isinstance(item, dict): add_text(f"• {_clean(item.get('priority'))} — {_clean(item.get('text'))}{_refs(item)}")
        else: add_text("• " + _clean(item))
    add_heading(data["labels"]["sources"])
    for source in data["sources"]: add_text(f"[{_clean(source.get('source_id')) or 'S?'}] {_clean(source.get('title'))} — {_clean(source.get('url'))}")
    story.extend([Spacer(1, 8), Paragraph(escape(data["labels"]["audit_note"]), meta), Paragraph(escape(f"Task: {data['task_id']}"), meta)])
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm).build(story)


def create_human_report(*, task_id: str, title: str, request: str, handoff_path: str | Path, artifact_root: str | Path, language: str, llm: Any | None = None) -> HumanReportBundle:
    path = Path(handoff_path)
    if not path.is_file():
        raise FileNotFoundError(f"Research handoff not found: {path}")
    handoff = json.loads(path.read_text(encoding="utf-8"))
    data = build_report_data(task_id, title, request, handoff, language)
    report_llm = llm or _runtime_report_llm()
    if report_llm is not None:
        data = compose_expert_report(data, handoff, report_llm)
    elif os.getenv("THREE_AGENT_CONFIG", "").strip():
        raise RuntimeError("Human report model is unavailable in installed runtime; refusing to publish an unlocalized fallback.")
    if report_llm is not None and not _language_ok(_language_body(data), data["language"]):
        raise ValueError("Human report rejected by final target-language validation.")
    markdown = render_markdown(data)
    folder = Path(artifact_root) / "reports" / datetime.now(TZ).strftime("%Y-%m-%d"); folder.mkdir(parents=True, exist_ok=True)
    stem = f"{task_id}_report"; md_path, docx_path, pdf_path = folder / f"{stem}.md", folder / f"{stem}.docx", folder / f"{stem}.pdf"
    md_path.write_text(markdown, encoding="utf-8"); paths, warnings = [str(md_path)], []
    try: _write_docx(data, docx_path); paths.append(str(docx_path))
    except Exception as exc: warnings.append(f"DOCX export failed: {type(exc).__name__}: {exc}")
    try: _write_pdf(data, pdf_path); paths.append(str(pdf_path))
    except Exception as exc: warnings.append(f"PDF export failed: {type(exc).__name__}: {exc}")
    return HumanReportBundle(markdown=markdown, paths=tuple(paths), warnings=tuple(warnings))
